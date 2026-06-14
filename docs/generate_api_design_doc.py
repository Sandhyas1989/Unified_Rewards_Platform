# -*- coding: utf-8 -*-
"""Generates the API Design Document for the as-built Unified Rewards microservices platform."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.dirname(os.path.abspath(__file__))
ACCENT = RGBColor(0x31, 0x2E, 0x81)
HEADER_FILL = "312E81"
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def shade(c, hexf):
    tcPr = c._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexf); tcPr.append(sh)


def h1(t):
    p = doc.add_heading(t, 1)
    for r in p.runs: r.font.color.rgb = ACCENT


def h2(t):
    p = doc.add_heading(t, 2)
    for r in p.runs: r.font.color.rgb = ACCENT


def para(t, italic=False, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic; r.bold = bold; return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0] + " "); r.bold = True; p.add_run(it[1])
        else:
            p.add_run(it)


def table(headers, rows, widths=None, fs=8):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, ht in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(ht); r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(fs)
        shade(c, HEADER_FILL)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(fs)
    if widths:
        for i, w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width = Inches(w)
    doc.add_paragraph()


def figure(fname, caption, width=6.6):
    doc.add_picture(os.path.join(OUT, fname), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


EP = ["Method", "Path (via gateway)", "Auth / Role", "Request body", "Success response"]
W = [0.7, 2.0, 1.2, 1.6, 1.5]

# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Unified Rewards Platform"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = ACCENT
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("API Design Document — Microservices (as built)"); r.font.size = Pt(15); r.italic = True
mm = doc.add_paragraph(); mm.alignment = WD_ALIGN_PARAGRAPH.CENTER
mm.add_run("Gateway + 7 backend microservices. Contracts, authentication, routing, and cross-service flows "
           "of the implemented system, with the production (Azure) mapping noted throughout.").font.size = Pt(11)
doc.add_paragraph()
nn = doc.add_paragraph(); nn.alignment = WD_ALIGN_PARAGRAPH.CENTER
nn.add_run("Source of truth: docs/Requirement_Document. Cognizant — Upgrade to Architect, Phase 3.").italic = True
doc.add_page_break()

h1("1. Purpose & scope")
para("This document specifies the HTTP API of the Unified Rewards microservices platform exactly as "
     "implemented and verified: one API gateway and seven backend services, each owning its data. All "
     "endpoints are reached through the gateway. Where the local implementation differs from the intended "
     "production deployment, the Azure mapping is called out.")

h1("2. Architecture overview")
figure("as_built_microservices_components.png",
       "Figure 1 — As-built component view: the gateway routes to seven services, each owning its database; "
       "services collaborate over HTTP (the local stand-in for Azure Service Bus).")

h1("3. API conventions")
bullets([
    ("Single entry point.", " All client traffic goes through the gateway (local http://localhost:5080 → Azure API Management in production). Service ports (5101–5107) are internal."),
    ("Authentication.", " Bearer JWT in the Authorization header. The token is issued by the Employee Profile service’s /api/auth/login (→ Microsoft Entra ID in production) and validated by every service."),
    ("Multi-tenancy.", " The token carries a tenant_id claim; every service scopes all reads/writes to that tenant automatically."),
    ("Authorisation.", " Role-based (Employee, Manager, HrAdmin, Finance) via the token’s role claim; unauthorised → 401, forbidden → 403."),
    ("Pagination.", " List endpoints accept page & pageSize (default 25, max 200) and return { items, page, pageSize, totalCount }. No endpoint returns an unbounded set."),
    ("Errors.", " JSON problem objects; standard status codes (400 invalid, 401 unauthenticated, 403 forbidden, 404 not found, 409 conflict, 502 upstream)."),
    ("Content type.", " application/json, except receipt upload (multipart/form-data) and report export (.xlsx)."),
])

h1("4. Gateway routing")
table(["Path prefix", "Service", "Local port"],
      [["/api/auth, /api/employees", "Employee Profile", "5101"],
       ["/api/benefit-plans, /api/enrollments", "Benefits Catalogue", "5102"],
       ["/api/compensation", "Compensation Rules", "5103"],
       ["/api/claims", "Reimbursement Workflow", "5104"],
       ["/api/documents", "Document & Receipt Processing", "5105"],
       ["/api/settlements, /api/payslips", "Payroll Integration", "5106"],
       ["/api/reports", "Reporting & Compliance", "5107"]],
      widths=[2.6, 2.6, 1.0], fs=9)

h1("5. Endpoint catalogue")

h2("5.1 Employee Profile  (identity, users)")
table(EP, [
    ["POST", "/api/auth/login", "Anonymous", "{ email, password }", "200 { token, expiresAtUtc, user }"],
    ["POST", "/api/employees", "HrAdmin", "{ fullName, email, password, grade, dateOfJoining, managerId? }", "201 UserDto"],
    ["GET", "/api/employees?role=&page=&pageSize=", "Authenticated", "—", "200 Paged<UserDto>"],
    ["GET", "/api/employees/{id}", "Authenticated", "—", "200 UserDto / 404"],
], widths=W)

h2("5.2 Benefits Catalogue  (plans, enrolments)")
table(EP, [
    ["GET", "/api/benefit-plans?category=&activeOnly=&page=&pageSize=", "Authenticated", "—", "200 Paged<BenefitPlanDto>"],
    ["POST", "/api/benefit-plans", "HrAdmin", "{ name, description, category, monthlyCost }", "201 / 409 (dup name)"],
    ["POST", "/api/enrollments", "Employee, Manager", "{ benefitPlanId, coverageStartDate }", "201 / 400 / 409 (dup)"],
    ["DELETE", "/api/enrollments/{id}", "Employee, Manager (owner)", "—", "204 / 404"],
    ["GET", "/api/enrollments/me", "Authenticated", "—", "200 [BenefitEnrollmentDto]"],
], widths=W)

h2("5.3 Compensation Rules  (rule-engine output)")
table(EP, [
    ["POST", "/api/compensation", "HrAdmin", "{ employeeId, grade, annualBasic, effectiveFrom }", "201 CompensationStructureDto (NRules)"],
    ["POST", "/api/compensation/{id}/approve", "HrAdmin, Finance", "—", "200 / 404"],
    ["GET", "/api/compensation?employeeId=", "Authenticated", "—", "200 [CompensationStructureDto]"],
    ["GET", "/api/compensation/me", "Authenticated", "—", "200 [CompensationStructureDto]"],
    ["GET", "/api/compensation/{id}", "Authenticated", "—", "200 / 404"],
], widths=W)

h2("5.4 Reimbursement Workflow  (claims state machine + saga)")
table(EP, [
    ["POST", "/api/claims", "Employee, Manager", "{ type, amount, description }", "201 ClaimDto (Submitted)"],
    ["POST", "/api/claims/{id}/approve", "Manager, Finance, HrAdmin", "{ notes? }", "200 / 409 (illegal transition)"],
    ["POST", "/api/claims/{id}/reject", "Manager, Finance, HrAdmin", "{ notes? }", "200 / 409"],
    ["POST", "/api/claims/{id}/settle", "Finance", "—", "200 Settled (orchestrates Payroll) / 409 / 502"],
    ["GET", "/api/claims?status=&page=&pageSize=", "Manager, Finance, HrAdmin", "—", "200 Paged<ClaimDto>"],
    ["GET", "/api/claims/me", "Authenticated", "—", "200 [ClaimDto]"],
    ["GET", "/api/claims/{id}", "Manager, Finance, HrAdmin", "—", "200 / 404"],
], widths=W)

h2("5.5 Document & Receipt Processing  (storage + OCR)")
table(EP, [
    ["POST", "/api/documents", "Authenticated", "multipart: ClaimId, File", "201 DocumentDto (stored + OCR-scanned)"],
    ["GET", "/api/documents?claimId=&page=&pageSize=", "Authenticated", "—", "200 Paged<DocumentDto>"],
    ["GET", "/api/documents/{id}", "Authenticated", "—", "200 / 404"],
    ["GET", "/api/documents/{id}/file", "Authenticated", "—", "200 (binary file)"],
], widths=W)

h2("5.6 Payroll Integration  (settlements, payslips)")
table(EP, [
    ["POST", "/api/settlements", "Finance", "{ employeeId, amount }", "202 SettlementDto (async, Polly-resilient)"],
    ["GET", "/api/settlements/{id}", "Authenticated", "—", "200 SettlementDto (poll for status)"],
    ["GET", "/api/settlements?page=&pageSize=", "Authenticated", "—", "200 Paged<SettlementDto>"],
    ["POST", "/api/payslips", "Finance, HrAdmin", "{ employeeId, year, month, grossMonthly, totalDeductionsMonthly, netMonthly }", "201 (idempotent per period)"],
    ["GET", "/api/payslips?employeeId=&page=&pageSize=", "Authenticated", "—", "200 Paged<PayslipDto>"],
], widths=W)

h2("5.7 Reporting & Compliance  (read model)")
table(EP, [
    ["GET", "/api/reports/dashboard", "HrAdmin, Finance", "—", "200 DashboardDto (aggregated)"],
    ["GET", "/api/reports/claims/export", "HrAdmin, Finance", "—", "200 .xlsx workbook"],
], widths=W)

h1("6. Cross-service flows")
para("The reimbursement saga shows how services collaborate. Locally this is direct service-to-service "
     "HTTP (orchestration); in production the settle step publishes an event to Azure Service Bus "
     "(choreography). Reporting & Compliance aggregates live from Reimbursement and Payroll (a derived "
     "read model that consumes Service Bus events in production).")
figure("reimbursement_saga_sequence.png",
       "Figure 2 — Reimbursement saga: submit → approve → settle, where ‘settle’ orchestrates the Payroll "
       "service (request, asynchronous Polly-resilient processing, poll), then closes the claim.")

h1("7. Data ownership (database-per-service)")
para("Each service is the only owner/writer of its data; others access it through the API or via events. "
     "No shared database.")
table(["Service", "Owns (entities / tables)", "Store (local → prod)"],
      [["Employee Profile", "Users (TPH: Employee/Manager/HrAdmin/Finance)", "SQLite → Azure SQL"],
       ["Benefits Catalogue", "BenefitPlans, BenefitEnrollments", "SQLite → Azure SQL"],
       ["Compensation Rules", "CompensationStructures, CompensationComponents", "SQLite → Azure SQL"],
       ["Reimbursement Workflow", "Claims, ClaimTransitions (history)", "SQLite → Azure SQL"],
       ["Document & Receipt", "Documents (+ files via storage seam)", "SQLite + filesystem → Azure SQL + Blob"],
       ["Payroll Integration", "Settlements, Payslips", "SQLite → Azure SQL"],
       ["Reporting & Compliance", "(derived read model — no transactional store)", "in-memory → event-sourced store"]],
      widths=[1.8, 3.2, 2.0], fs=9)

h1("8. Production (Azure) mapping")
table(["Concern", "Local (as built)", "Production (Azure)"],
      [["Gateway", "YARP reverse proxy", "Azure API Management"],
       ["Identity", "JWT issued by Employee Profile (shared key)", "Microsoft Entra ID (OIDC); managed identity for service-to-service"],
       ["Service-to-service", "Direct HTTP (orchestration)", "Azure Service Bus events (choreography) + outbox"],
       ["Datastores", "SQLite per service", "Azure SQL / Cosmos DB per service"],
       ["Documents", "Local filesystem", "Azure Blob Storage"],
       ["OCR", "Stub (reads text receipts)", "Azure AI Document Intelligence"],
       ["Hosting", "dotnet run processes", "Containers on AKS (HPA per service)"],
       ["Secrets / observability", "config files / console", "Key Vault / Application Insights"]],
      widths=[1.6, 2.6, 2.8], fs=8.5)

doc.save(os.path.join(OUT, "Unified_Rewards_Platform_API_Design.docx"))
print("SAVED API design doc")
