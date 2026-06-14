# -*- coding: utf-8 -*-
"""Generates the Microservices Migration & Decomposition Plan (Word doc) + a target
microservices architecture diagram, grounded in docs/Requirement_Document (the 7 named services)."""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.dirname(os.path.abspath(__file__))
INK = "#0f172a"
ACCENT = RGBColor(0x31, 0x2E, 0x81)
HEADER_FILL = "312E81"

# ---------------------------------------------------------------- diagram
PAL = {
    "user": ("#EEF2FF", "#6366F1"), "edge": ("#E0E7FF", "#4F46E5"), "fe": ("#DBEAFE", "#2563EB"),
    "svc": ("#DCFCE7", "#16A34A"), "data": ("#FEF3C7", "#D97706"), "bus": ("#FCE7F3", "#DB2777"),
    "cross": ("#E2E8F0", "#475569"), "band": ("#F8FAFC", "#CBD5E1"),
}


def b(ax, x, y, w, h, t, k, fs=8, bold=False):
    f, e = PAL[k]
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.1",
                                lw=1.3, edgecolor=e, facecolor=f))
    ax.text(x + w / 2, y + h / 2, t, ha="center", va="center", fontsize=fs, color=INK,
            weight="bold" if bold else "normal")


def band(ax, x, y, w, h, label):
    f, e = PAL["band"]
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.15",
                                lw=1.4, edgecolor=e, facecolor=f, linestyle="--"))
    ax.text(x + 0.7, y + h - 0.8, label, ha="left", va="top", fontsize=9, color="#334155", weight="bold")


def arr(ax, x1, y1, x2, y2, t=None):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color="#475569", lw=1.4, shrinkA=2, shrinkB=2))
    if t:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2, t, ha="center", va="center", fontsize=7,
                color="#475569", backgroundcolor="white")


fig, ax = plt.subplots(figsize=(14, 11))
ax.set_xlim(0, 100); ax.set_ylim(0, 100); ax.axis("off")
ax.text(50, 98.5, "Target Microservices Architecture on Azure (per region — multi-tenant, multi-country)",
        ha="center", fontsize=14, weight="bold", color=INK)

b(ax, 33, 92, 34, 4.5, "Users — multi-tenant, global (Employee · Manager · HR · Finance)", "user", bold=True)
arr(ax, 50, 92, 50, 89.5)
b(ax, 28, 85, 44, 4.5, "Azure Front Door + CDN + WAF (TLS, edge cache, routing by region/tenant)", "edge", bold=True)
arr(ax, 40, 85, 26, 79)
b(ax, 5, 73.5, 40, 5, "Azure Static Web Apps — 4 micro-frontends\n(Shell + Employee/Manager/HR/Finance, Module Federation)", "fe", fs=7.5)
arr(ax, 62, 85, 74, 79)
b(ax, 55, 73, 40, 6, "Azure API Management (gateway)\nrouting · validate Entra JWT · per-tenant rate-limit · versioning", "edge", fs=7.5, bold=True)
arr(ax, 75, 73, 60, 67)
arr(ax, 25, 73.5, 45, 67)

# AKS band with 7 services
band(ax, 4, 40, 74, 26, "Azure Kubernetes Service (AKS) — one deployment + HPA autoscale per service; mTLS service mesh; Key Vault CSI")
svcs = ["Employee\nProfile", "Benefits\nCatalogue", "Compensation\nRules Engine", "Reimbursement\nWorkflow",
        "Document &\nReceipt Proc.", "Payroll\nIntegration", "Reporting &\nCompliance"]
for i, s in enumerate(svcs):
    col = i % 4; row = i // 4
    x = 7 + col * 17.5
    y = 53 if row == 0 else 42
    b(ax, x, y, 15.5, 10, s + "\nService", "svc", fs=7.3, bold=True)

# Event bus
b(ax, 4, 33, 74, 4.5, "Azure Service Bus — event backbone (ClaimApproved · DocProcessed · SettlementRequested · CompApproved · audit events)", "bus", fs=7.5, bold=True)
arr(ax, 41, 40, 41, 37.5)

# Data row — database per service
band(ax, 4, 14, 74, 17, "Data ownership — database-per-service (no shared DB)")
b(ax, 6, 23.5, 16, 6, "Azure SQL\nper service", "data", fs=7.5, bold=True)
b(ax, 24, 23.5, 16, 6, "Azure Cosmos DB\n(catalogue / docs meta)", "data", fs=7)
b(ax, 42, 23.5, 16, 6, "Azure Cache\nfor Redis", "data", fs=7.5)
b(ax, 60, 23.5, 16, 6, "Azure Blob\n(receipts/docs)", "data", fs=7.5)
b(ax, 24, 15.5, 34, 6, "Reporting read-store / analytics (event-sourced from Service Bus)", "data", fs=7)

# Cross-cutting column on right
band(ax, 80, 14, 16, 65, "Cross-cutting")
b(ax, 81.5, 70, 13, 7, "Microsoft\nEntra ID\nSSO·MFA·tenant", "cross", fs=7)
b(ax, 81.5, 60, 13, 7, "Azure\nKey Vault\n(secrets)", "cross", fs=7)
b(ax, 81.5, 50, 13, 7, "App Insights\n+ OpenTelemetry\n(tracing)", "cross", fs=7)
b(ax, 81.5, 40, 13, 7, "Azure\nContainer\nRegistry", "cross", fs=7)
b(ax, 81.5, 27, 13, 9, "GitHub Actions\nCI/CD per\nservice →\nACR → AKS", "cross", fs=7)
b(ax, 81.5, 16, 13, 8, "Per-region\ndeploy for\ndata residency\n(compliance)", "cross", fs=7)

fig.tight_layout()
fig.savefig(os.path.join(OUT, "target_microservices_architecture.png"), dpi=150, bbox_inches="tight")
plt.close(fig)

# ---------------------------------------------------------------- document
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def shade(c, hexf):
    tcPr = c._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexf); tcPr.append(sh)


def h1(t):
    p = doc.add_heading(t, 1)
    for r in p.runs: r.font.color.rgb = ACCENT


def h2(t):
    p = doc.add_heading(t, 2)
    for r in p.runs: r.font.color.rgb = ACCENT


def para(t, italic=False, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic; r.bold = bold; return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0] + " "); r.bold = True; p.add_run(it[1])
        else:
            p.add_run(it)


def table(headers, rows, widths=None, fs=8.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, ht in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(ht); r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(fs)
        shade(c, HEADER_FILL)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(fs)
    if widths:
        for i, w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width = Inches(w)
    doc.add_paragraph()


def figure(fname, caption, width=6.6):
    doc.add_picture(os.path.join(OUT, fname), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Unified Rewards Platform"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = ACCENT
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Microservices Migration & Decomposition Plan"); r.font.size = Pt(15); r.italic = True
mm = doc.add_paragraph(); mm.alignment = WD_ALIGN_PARAGRAPH.CENTER
mm.add_run("Evolving the modular monolith into the requirement-mandated 7-service architecture on Azure — "
           "multi-tenant, multi-country, AKS-hosted.").font.size = Pt(11)
doc.add_paragraph()
nn = doc.add_paragraph(); nn.alignment = WD_ALIGN_PARAGRAPH.CENTER
nn.add_run("Source of truth: docs/Requirement_Document. Cognizant — Upgrade to Architect, Phase 3.").italic = True
doc.add_page_break()

h1("1. Purpose & scope")
para("The Requirement_Document mandates a Backend-Microservices architecture (seven named services) plus "
     "micro-frontends, for a Large, multi-tenant HR-Tech Enterprise SaaS on Azure. The platform today is a "
     "modular monolith (functionally complete; micro-frontends already delivered). This plan defines the "
     "target service decomposition, data ownership, communication, API gateway, multi-tenancy, multi-country "
     "compliance, AKS deployment, and a safe, incremental migration path.")
para("Key principle: the monolith was deliberately built with Clean Architecture and clear module seams, so "
     "each module maps almost 1:1 to a target service and can be carved out using the strangler pattern — an "
     "evolution, not a rewrite.")

h1("2. Service decomposition — mapping the 7 modules to the 7 services")
para("The requirement names seven services. The mapping below shows two decomposition decisions: the "
     "monolith’s combined Claims & Documents module splits into two services (as the requirement separates "
     "them), and Promotions — which the requirement does not list as a standalone service — folds into the "
     "Employee Profile / Compensation domains as a capability that emits events.")
table(["Current monolith module", "Target microservice (per requirement)", "Decomposition note"],
      [["User Management", "Employee Profile Service", "AuthN/AuthZ moves to Entra ID; service owns profile, org, grade, promotion history"],
       ["Benefits", "Benefits Catalogue Service", "Plans + enrolments; mostly catalogue/CRUD"],
       ["Compensation (NRules)", "Compensation Rules Engine", "Comp structures + country-specific rule sets"],
       ["Claims & Documents", "Reimbursement Workflow Service", "Owns the claim state machine & approvals"],
       ["Claims & Documents", "Document & Receipt Processing Service", "SPLIT OUT: receipt storage + OCR (different scaling profile)"],
       ["Payroll", "Payroll Integration Service", "External payroll integration + settlements (Polly resilience)"],
       ["Promotions", "→ Employee Profile / Compensation", "No standalone service in the requirement; modelled as a capability + events"],
       ["Reporting & Audit", "Reporting & Compliance Service", "Event-sourced read store; compliance reporting + audit"]],
      widths=[1.8, 2.2, 2.8])

h2("Per-service responsibilities & dependencies")
table(["Service", "Owns (data)", "Sync calls (via APIM)", "Async events (Service Bus)"],
      [["Employee Profile", "Employees, org, grade", "—", "publishes EmployeeOnboarded, GradeChanged"],
       ["Benefits Catalogue", "Plans, enrolments", "Employee Profile (validate)", "publishes EnrolmentChanged"],
       ["Compensation Rules", "Comp structures, rules", "Employee Profile", "publishes CompensationApproved; consumes GradeChanged"],
       ["Reimbursement Workflow", "Claims, approvals", "Employee Profile", "publishes ClaimSubmitted/Approved/Settled; consumes DocProcessed, SettlementCompleted"],
       ["Document & Receipt", "Receipts (Blob), OCR results", "—", "consumes ClaimSubmitted; publishes DocProcessed"],
       ["Payroll Integration", "Settlements, payslips", "—", "consumes ClaimApproved; publishes SettlementCompleted"],
       ["Reporting & Compliance", "Read models, audit", "(reads replicas)", "consumes ALL events"]],
      widths=[1.5, 1.6, 1.5, 2.2], fs=8)

h1("3. Data ownership — database-per-service")
bullets([
    ("No shared database.", " Each service owns its data and is the only writer; others get data via its API or via events. This is what enables independent deployment and scaling."),
    ("Provider per fit.", " Relational, transactional services (Reimbursement, Compensation, Payroll, Employee) use Azure SQL; catalogue/读-heavy or document-metadata services may use Azure Cosmos DB; documents in Blob; hot reads in Redis."),
    ("Splitting the current schema.", " The monolith’s single schema is partitioned along module boundaries; foreign keys that cross a boundary become API/event references (e.g., a claim stores EmployeeId, not a FK to a Users table it no longer owns)."),
    ("Consistency across services.", " Use the Saga pattern (orchestrated by Reimbursement Workflow) with the Transactional Outbox so events are never lost; accept eventual consistency between services; compensating actions on failure."),
    ("Reporting.", " Reporting & Compliance builds its own read store by subscribing to events (event-sourced/CQRS read model) — it never queries other services’ databases directly."),
])

h1("4. Inter-service communication")
bullets([
    ("Synchronous (request/response).", " REST over HTTPS through API Management for queries that need an immediate answer (e.g., validate an employee). Kept to a minimum to avoid coupling and latency chains."),
    ("Asynchronous (events).", " Azure Service Bus topics/queues for state changes — the default. Services react to events rather than calling each other, which decouples them and absorbs load spikes."),
    ("The reimbursement saga.", " Submit → (event) Document & Receipt OCRs the receipt → Reimbursement Workflow advances on DocProcessed → on approval, (event) Payroll settles → SettlementCompleted closes the claim. Each step is independently retriable and idempotent."),
])
para("Event backbone (illustrative): EmployeeOnboarded, GradeChanged, EnrolmentChanged, CompensationApproved, "
     "ClaimSubmitted, ClaimApproved, ClaimRejected, DocProcessed, SettlementRequested, SettlementCompleted, "
     "plus audit/compliance events consumed by Reporting & Compliance.", italic=True)

h1("5. API Management (APIM) — the gateway")
bullets([
    ("Single entry point", " for all micro-frontends and external integrations; routes to the right service."),
    ("Security", " — validates the Entra ID JWT (incl. the tenant claim), enforces scopes/roles, terminates TLS."),
    ("Protection", " — per-tenant rate limiting/throttling and quotas; protects services from spikes and noisy tenants."),
    ("Versioning & lifecycle", " — API versions, revisions, and a developer portal for enterprise integration partners."),
    ("Aggregation (BFF)", " — optional backend-for-frontend composition so a portal screen needs one call, not many."),
])

h1("6. Multi-tenancy (the SaaS dimension)")
para("Each customer enterprise is a tenant. Tenant identity flows from Entra ID as a tenant claim in the JWT, "
     "is validated at APIM, and is propagated to every service and onto every data operation.")
table(["Isolation model", "What it means", "Use for"],
      [["Pooled (shared DB + TenantId)", "One database, every row tagged with TenantId, enforced by a global query filter / row-level security", "Default — most tenants; best density & cost"],
       ["Bridge (schema-per-tenant)", "Shared server, separate schema per tenant", "Mid-size tenants needing logical separation"],
       ["Siloed (DB-per-tenant)", "Dedicated database (and region) per tenant", "Large or regulated tenants; data-residency needs"]],
      widths=[2.0, 3.4, 1.4])
bullets([
    ("Recommended: a hybrid ‘pool + silo’.", " Pooled by default; promote specific tenants to siloed/regional databases when contracts or compliance demand it — without code changes (tenant-to-store routing is configuration)."),
    ("Tenant context everywhere.", " EF Core global query filters guarantee no query can cross tenant boundaries; APIM rejects tokens without a valid tenant; audit records the tenant."),
    ("Tenant onboarding.", " A provisioning process creates the tenant, its store/route, default catalogue, and admin user."),
])

h1("7. Multi-country compliance")
bullets([
    ("Data residency.", " Deploy per-region (AKS cluster + data stores per geography); route each tenant to its home region via Front Door/APIM so data stays in-country."),
    ("Country-specific business rules.", " The Compensation Rules Engine and Reimbursement Workflow load country rule sets (tax, statutory allowances, claim limits) — a natural fit for the existing NRules engine, with one rule set per country/tenant."),
    ("Privacy (GDPR & local laws).", " Consent capture, data-minimisation, right-to-erasure/portability workflows, and field-level encryption for sensitive pay/PII."),
    ("Auditability.", " Reporting & Compliance maintains an immutable audit/compliance trail (from events) with per-country retention policies."),
    ("Localisation.", " Currency, number/date formats, language (i18n) per country in the micro-frontends."),
])

h1("8. Deployment on AKS")
bullets([
    ("One deployment + HPA per service", " — each service scales independently on CPU/memory or queue length; no more scaling the whole app to scale one part."),
    ("Namespaces per environment", " (dev/test/prod); ingress controller behind APIM/Front Door."),
    ("Service mesh (optional, e.g. Istio/Linkerd)", " for mutual-TLS between services, traffic shifting (canary), and uniform telemetry."),
    ("Secrets via Key Vault CSI driver + Workload Identity", " — no secrets in images or config; pods get managed-identity access."),
    ("Images in Azure Container Registry", "; GitHub Actions pipeline per service builds, scans, and deploys (Helm/Kustomize) with blue-green/canary."),
    ("Per-region clusters", " for residency and lower latency; geo-failover for DR."),
])
figure("target_microservices_architecture.png",
       "Figure 1 — Target microservices architecture on Azure (per region). Stateless services on AKS behind "
       "APIM, communicating via Service Bus events, each owning its data, with Entra ID, Key Vault and "
       "Application Insights as cross-cutting platform services.")

h1("9. Cross-cutting concerns")
bullets([
    ("Identity", " — Microsoft Entra ID (OIDC/OAuth2, RS256, MFA, SSO) issues tenant- and role-aware tokens; services trust the gateway-validated token."),
    ("Observability", " — Application Insights + OpenTelemetry distributed tracing with a correlation id propagated across sync calls and async events; per-service dashboards and SLO alerts."),
    ("Resilience", " — Polly retries/circuit-breakers/timeouts on every external and inter-service call (already proven on Payroll); idempotent consumers."),
    ("Shared contracts", " — event and DTO contracts versioned in a shared package; consumer-driven contract tests prevent breaking changes."),
    ("Config & secrets", " — Azure App Configuration + Key Vault; feature flags for safe rollout."),
])

h1("10. Migration strategy & sequencing (Strangler pattern)")
para("The monolith keeps running and serving traffic while services are carved out one at a time behind APIM; "
     "each extracted capability is routed to its new service and the old code path retired. No big-bang cutover.")
table(["Phase", "Focus", "Key work"],
      [["0. Foundations", "Platform", "Entra ID, APIM, Service Bus, AKS, ACR, observability, shared-contracts lib, CI/CD, tenant context + multi-tenant data filter"],
       ["1. Front the monolith", "Gateway", "Put APIM in front of the existing monolith (no behaviour change); switch auth to Entra ID; add tenant claim end-to-end"],
       ["2. Extract leaf services", "Low-coupling first", "Document & Receipt Processing, then Payroll Integration (already behind seams, async) — own data store + events; route via APIM"],
       ["3. Extract domain services", "Workflow & catalogue", "Benefits Catalogue, Compensation Rules Engine, Reimbursement Workflow (saga over Document + Payroll)"],
       ["4. Extract core + reporting", "Core & read-side", "Employee Profile Service; Reporting & Compliance built as an event-sourced read store consuming all events"],
       ["5. Retire monolith", "Cutover & harden", "Remove residual monolith code; full multi-region + compliance + DR; load/chaos/security testing"]],
      widths=[1.3, 1.4, 4.1], fs=8.5)
para("Extraction order rationale: start with the most independent, differently-scaling capability (Document/OCR), "
     "then the already-seamed async Payroll, building the event backbone early; tackle the highly-referenced "
     "Employee Profile later when the event patterns are proven; Reporting consumes events throughout.", italic=True)

h1("11. Risks & mitigations")
table(["Risk", "Mitigation"],
      [["Distributed data consistency", "Saga + Transactional Outbox; idempotency; design for eventual consistency"],
       ["Operational complexity (many services)", "IaC, service mesh, strong observability, platform team & golden paths"],
       ["Latency from chatty sync calls", "Prefer events; cache; BFF aggregation; keep sync calls shallow"],
       ["Harder testing", "Consumer-driven contract tests; per-service test suites; staging with prod-like data"],
       ["Cost of AKS + per-service stores", "Right-size, autoscale to zero where possible, pooled multi-tenancy by default"],
       ["Migration regressions", "Strangler (incremental), feature flags, canary, parallel-run & compare"]],
      widths=[2.6, 4.2])

h1("12. How this satisfies the Requirement_Document")
bullets([
    ("Backend Microservices (7):", " delivered exactly as named (with Claims/Documents split and Promotions folded)."),
    ("Micro Frontends (4):", " already delivered via React Module Federation; retained as-is."),
    ("Rule engines vs hard-coded:", " Compensation Rules Engine (NRules) + country rule sets."),
    ("Multi-country compliance:", " per-region residency, country rule sets, GDPR workflows, audit."),
    ("Enterprise integrations:", " APIM developer portal + Payroll Integration Service + event backbone."),
    ("Data privacy & access segregation:", " Entra ID + RBAC + per-tenant isolation + field-level encryption."),
    ("Scalable workflow design:", " event-driven saga on Service Bus, each service autoscaled on AKS."),
])

doc.save(os.path.join(OUT, "Unified_Rewards_Platform_Microservices_Migration_Plan.docx"))
print("SAVED doc + diagram")
