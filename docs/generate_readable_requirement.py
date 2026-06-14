# -*- coding: utf-8 -*-
"""Readable .docx of docs/Requirement_Document.doc. The original is IRM rights-protected, which
blocks programmatic Save-As/Export of its 1 embedded image. ALL TEXT is reproduced faithfully
(verified 2,520 chars via Word COM Content.Text); the single image is marked with a placeholder."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.dirname(os.path.abspath(__file__))
ACCENT = RGBColor(0x31, 0x2E, 0x81)
AMBER = RGBColor(0xB4, 0x53, 0x09)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def h1(t):
    p = doc.add_heading(t, 1)
    for r in p.runs: r.font.color.rgb = ACCENT


def h2(t):
    p = doc.add_heading(t, 2)
    for r in p.runs: r.font.color.rgb = ACCENT


def para(t, italic=False, bold=False, color=None):
    par = doc.add_paragraph(); r = par.add_run(t); r.italic = italic; r.bold = bold
    if color: r.font.color.rgb = color
    return par


def bul(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def placeholder(text):
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(text); r.bold = True; r.font.color.rgb = AMBER; r.font.size = Pt(10)


# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Enterprise Employee Benefits & Compensation Management Platform")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ACCENT
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Requirement Document — Readable Edition"); r.italic = True; r.font.size = Pt(14)
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run("Upgrade to Architect · April 2026").font.size = Pt(11)

# Provenance / completeness note
doc.add_paragraph()
note = doc.add_paragraph()
rr = note.add_run("About this edition: ")
rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = AMBER
rr2 = note.add_run("Every word of the original docs/Requirement_Document.doc is reproduced below "
                   "(verified: 2,520 characters). The original is rights-protected (IRM), which prevents "
                   "software from extracting its one embedded image; that picture is marked with a "
                   "placeholder. To include the actual image, open the original in Word (your account has "
                   "the rights) and use File ▸ Save As ▸ Word Document (.docx) — that copy will contain the image.")
rr2.italic = True; rr2.font.size = Pt(9); rr2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
doc.add_paragraph()

h1("1. Problem Statement")
para("A configurable platform for enterprises to manage employee benefits, reimbursements, bonuses, "
     "and flexible compensation structures.")

h1("2. Context")
bul([
    "Domain: HR Tech / Enterprise SaaS",
    "Scale: Large",
    "Cloud: AWS or Azure",
    "Technology Stack: AWS stack or Azure stack with any programming language of your choice, "
    "leveraging any of the Gen AI code companion tools.",
])

h1("3. Features / Requirements — Key Capabilities")
bul([
    "Benefits catalogue (insurance, LTA, food cards)",
    "Flexible pay configuration",
    "Reimbursement workflow (claims, approvals)",
    "Payroll system integration",
    "Analytics & compliance reporting",
])

h1("4. Architecture Scope")
h2("Backend Microservices")
bul([
    "Employee Profile Service",
    "Benefits Catalogue Service",
    "Compensation Rules Engine",
    "Reimbursement Workflow Service",
    "Document & Receipt Processing Service",
    "Payroll Integration Service",
    "Reporting & Compliance Service",
])
h2("Micro Frontends")
bul([
    "Employee Self-Service Portal",
    "Manager Approval Portal",
    "HR Admin Console",
    "Finance & Audit Dashboard",
])

h1("5. Key Points to be Followed — Architect-Level Learning")
bul([
    "Rule engines vs hard-coded logic",
    "Multi-country compliance handling",
    "Enterprise integrations",
    "Data privacy and access segregation",
    "Scalable workflow design",
])

h1("6. Output to be Generated")
h2("Phase 1")
bul([
    "Application Architecture & Deployment Architecture",
    "High level design diagrams",
    "Low level design diagrams: class diagrams, sequence diagrams, ER diagrams, component diagram, "
    "deployment view — all mentioned",
    "Low level design document template",
    "API Design document: API views, API design, keys and everything as per API design document template",
    "Presentation",
])
h2("Phase 2: Execution (Code Companion Tools help)")
bul([
    "Code Output",
    "UI / Backend Demo",
    "Code Quality and Coverage Report",
    "All prompts (Failed as well Passed prompts)",
    "Screenshots",
])

h1("7. UI Details with Sample Snapshots")
h2("1) Employee Self-Service – Benefits Overview")
bul([
    "Health Insurance, Travel Allowance, Meal Card",
    "Coverage / allowance amounts",
    "Quick “Submit Claim” actions",
    "Compensation overview (Basic, Bonus, Total CTC)",
])
h2("2) Reimbursement / Claim Submission")
bul([
    "Claim type selection (Medical, Travel, LTA, Dental, etc.)",
    "Amount entry",
    "Receipt upload",
    "Optional notes",
    "Submit workflow",
])
h2("3) Claim History & Status Tracking")
bul([
    "Claim list with statuses: Approved, In Review, Rejected, Settled",
    "Amount visibility and auditability",
])
h2("4) HR / Admin – Benefits & Discounts Management")
bul([
    "Active promotions (e.g., Year-End Bonus, Diwali Voucher)",
    "Benefit configuration (coverage & limits)",
    "Edit / add benefits entry points",
])

doc.add_paragraph()
placeholder("[ The original document contains 1 embedded image (UI sample snapshot) here. "
            "It could not be auto-extracted because the source is rights-protected (IRM). "
            "Save the original as .docx in Word to include it. ]")

doc.save(os.path.join(OUT, "Requirement_Document_Readable.docx"))
print("SAVED:", os.path.join(OUT, "Requirement_Document_Readable.docx"))
