# -*- coding: utf-8 -*-
"""Generates the architecture & technology-decisions Word document for the
Unified Rewards Platform, written for a reader who is new to the system."""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_DIR = os.path.dirname(os.path.abspath(__file__))
doc = Document()

# ---------- base styling ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

ACCENT = RGBColor(0x31, 0x2E, 0x81)   # indigo
HEADER_FILL = "312E81"
SUBTLE_FILL = "EEF2FF"


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = ACCENT


def h3(text):
    doc.add_heading(text, level=3)


def p(text, italic=False, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = italic
    run.bold = bold
    return para


def bullets(items):
    for it in items:
        para = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = para.add_run(it[0] + " ")
            r.bold = True
            para.add_run(it[1])
        else:
            para.add_run(it)


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], HEADER_FILL)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def spacer():
    doc.add_paragraph()


def figure(filename, caption, width=6.3):
    path = os.path.join(DOC_DIR, filename)
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


# =====================================================================
# TITLE PAGE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Unified Rewards Platform")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Architecture & Technology Decisions — Explained from First Principles")
r.font.size = Pt(15)
r.italic = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    "A plain-language guide to every major choice in the system: what it is, "
    "why it was chosen, what was rejected, and how the local set-up maps to Microsoft Azure.\n"
    "Written for a reader who is brand-new to the application and to software architecture."
).font.size = Pt(11)

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.add_run("Cognizant — Upgrade to Architect, Phase 3 Case Study").italic = True

doc.add_page_break()

# =====================================================================
# 0. HOW TO READ
# =====================================================================
h1("0. How to read this document")
p("This document assumes you have never seen this application before and are not a software "
  "specialist. Every section is written so a newcomer can follow it. Wherever a technical term "
  "appears for the first time it is explained in plain words, and there is a Glossary at the very "
  "end that defines the recurring jargon.")
p("The structure is deliberate:")
bullets([
    ("Sections 1–2", "explain what the application does and the shape of the system, in everyday language."),
    ("Section 3", "explains how the application was divided into seven “modules” and why seven."),
    ("Sections 4–5", "go through every technology choice (backend, then frontend). Each choice follows the same pattern: What it is → Why we chose it → What we rejected and why."),
    ("Section 6", "is the local-versus-Azure map: what runs on a laptop during development and the matching Microsoft Azure service in production."),
    ("Section 7", "is a single consolidated “why not the alternatives” table for quick reference."),
    ("Sections 8–9", "cover cross-cutting concerns and the guiding philosophy."),
])

# =====================================================================
# 1. WHAT IS THE APP
# =====================================================================
h1("1. What is the Unified Rewards Platform?")
p("Imagine a mid-to-large company. Every employee receives more than just a salary — they get a "
  "whole package of “rewards”: health insurance and other benefits, a structured salary made up of "
  "many components, the ability to claim back expenses (travel, medical, etc.), monthly payslips, "
  "and the chance to be promoted. Today, in many companies, these things live in separate, "
  "disconnected systems and spreadsheets.")
p("The Unified Rewards Platform brings all of that into one place. It is a web application with two "
  "sides:")
bullets([
    ("The people who use it", "Employees, Managers, HR Administrators, and Finance staff. Each of these four groups sees a different “portal” tailored to what they are allowed to do."),
    ("The engine behind it", "A central server (the “backend”) that holds all the business rules and data, plus connections to outside services such as e-mail, a payroll system, document storage, and document scanning."),
])
p("A quick example that touches the whole system: an Employee submits an expense claim with a photo "
  "of a receipt → their Manager approves it → Finance settles it (which pushes a payment to the "
  "payroll system) → HR can see the whole action recorded in an audit log, and Finance can pull a "
  "report. Four different people, four different portals, one consistent platform underneath.")

# =====================================================================
# 2. BIG PICTURE
# =====================================================================
h1("2. The big picture (architecture at a glance)")
p("The system has three layers. Think of a restaurant: the dining room (what customers see), the "
  "kitchen (where the work happens), and the pantry/suppliers (where ingredients are stored and "
  "sourced).")
table(
    ["Layer", "Restaurant analogy", "In this system"],
    [
        ["Frontend (what users see)", "The dining room", "Four web “portals” (Employee, Manager, HR, Finance) running in the browser, built with React."],
        ["Backend (where work happens)", "The kitchen", "A .NET web API that enforces all business rules and decides who can do what."],
        ["Data & external services", "The pantry & suppliers", "A database plus outside services: e-mail, payroll, file storage, document scanning (OCR)."],
    ],
    widths=[1.8, 1.8, 3.2],
)
p("Two architectural ideas shape the whole platform, and both are explained in detail later:")
bullets([
    ("Modular monolith (backend)", "The backend is one deployable application, but internally it is split into clean, independent “modules” — one per business area. It gives the tidiness of separation without the heavy operational cost of many separate services."),
    ("Micro-frontends (frontend)", "The four portals are built and can run as independent applications, then stitched together at run-time into one seamless site using a technique called Module Federation."),
])
spacer()
figure("as_built_logical.png",
       "Figure 1 — As-built logical architecture: the Module Federation frontend (shell host + four portal "
       "remotes), the .NET Clean-Architecture modular monolith with its seven modules and layers, the "
       "integration ‘seams’, and the external resources behind them.")

# =====================================================================
# 3. WHY 7 MODULES
# =====================================================================
h1("3. How the seven modules were decided")
p("A “module” here means a self-contained slice of the system that handles one business "
  "capability end-to-end — its own data, its own rules, its own screens. The number seven was not "
  "arbitrary; it came from listing the real things the business needs to do and grouping them so "
  "that each group is cohesive (everything inside belongs together) and loosely coupled (groups "
  "depend on each other as little as possible). This is the core of Domain-Driven Design.")
p("Walking through the business capabilities, they naturally fall into these seven buckets:")
table(
    ["#", "Module", "The business question it answers", "Signature technology"],
    [
        ["1", "User Management", "Who are you, and what are you allowed to do?", "JWT login, roles, password hashing"],
        ["2", "Benefits", "Which perks can I enrol in, and manage?", "Straightforward data + rules"],
        ["3", "Compensation", "How is my salary structured into components?", "NRules business-rules engine"],
        ["4", "Claims & Documents", "Can I claim an expense and attach a receipt?", "State machine + OCR + file storage"],
        ["5", "Payroll", "Payslips and paying money out to staff", "Polly resilience + async background worker"],
        ["6", "Promotions", "Nominate and approve people for a higher grade", "Event-driven notifications"],
        ["7", "Reporting & Audit", "What happened, and show me the numbers", "LINQ aggregation + Excel export + audit trail"],
    ],
    widths=[0.3, 1.5, 3.0, 2.0],
)
h2("Why not fewer modules?")
p("We could have lumped, say, Compensation, Payroll, and Claims into one giant “Money” module. "
  "We did not, because each has a genuinely different shape: Compensation is rules-heavy, Payroll is "
  "about resilient communication with an outside system and background processing, and Claims is a "
  "step-by-step approval workflow. Forcing them together would mix unrelated concerns and make the "
  "code hard to change safely — the opposite of cohesion.")
h2("Why not more modules?")
p("We could have split further — e.g., “Documents” separate from “Claims”, or “Audit” "
  "separate from “Reporting”. We kept them together because they are almost always used together "
  "and share the same data and language. Splitting them would create extra boundaries that add "
  "friction without adding clarity. The guiding rule: a module should be big enough to be meaningful "
  "and small enough to be understood by one person in one sitting.")
p("Each module is built as a “vertical slice” — it reaches from the database all the way up to "
  "the web screens — and the very first module (User Management) was built as a reference template "
  "so the other six could follow exactly the same proven shape.")

# =====================================================================
# 4. BACKEND DECISIONS
# =====================================================================
h1("4. Backend technology decisions")
p("Each choice below follows the same three-part pattern so you can scan it quickly: what it is, "
  "why we chose it, and what we deliberately did not use.")

h2("4.1  Platform & language: .NET 8 and C#")
p("What it is: .NET is Microsoft’s framework for building applications; C# is its main "
  "programming language. .NET 8 is a long-term-support version.")
p("Why we chose it: the case study targets a Microsoft/Azure environment, and .NET is the "
  "first-class citizen there — best tooling, performance, hosting, and library support on Azure. "
  "It is strongly typed (the compiler catches many mistakes before the app ever runs), fast, and "
  "has a huge ecosystem of mature libraries for everything we needed.")
p("What we did not use, and why:")
bullets([
    ("Java/Spring Boot", "Equally capable, but a weaker fit for an Azure-first, Microsoft-centric brief; the team’s and the platform’s strengths point to .NET."),
    ("Node.js / NestJS", "Excellent for I/O-heavy services, but JavaScript’s looser typing and the preference for a single strongly-typed backend made .NET the safer choice for complex business rules."),
    ("Python/Django", "Great for data and scripting, but less aligned with enterprise Azure hosting and high-throughput typed APIs."),
])

h2("4.2  Overall shape: Clean Architecture as a modular monolith")
p("What it is: Clean Architecture organises code into concentric layers. The centre holds the "
  "business meaning; the outer rings hold technical details (database, web, external services). The "
  "key rule is that dependencies only point inward — the business core never depends on the "
  "database or the web framework. In this project the layers are: Domain (the business entities and "
  "rules), Application (the use-cases / what the system can do), Infrastructure (the technical "
  "implementations), and Api (the web entry point).")
p("Why we chose it: it keeps the valuable part of the software — the business logic — "
  "independent and testable, and it makes technical pieces replaceable. That replaceability is "
  "exactly what lets the same code run with a local file store on a laptop and Azure Blob Storage in "
  "the cloud, with no change to the business logic. “Modular monolith” means we get clean "
  "internal separation while still deploying one simple application.")
p("What we did not use, and why:")
bullets([
    ("Microservices (many separate deployable services)", "They solve problems of very large teams and independent scaling, but they add heavy operational complexity — networking, distributed transactions, monitoring, versioning. For a single team building a case study, that cost vastly outweighs the benefit. A modular monolith gives most of the organisational benefit now and can be split into services later if ever needed."),
    ("A traditional layered “N-tier” app without dependency inversion", "Simpler to start, but the business logic ends up tied to the database and framework, making change and testing painful."),
])

h2("4.3  Request handling: CQRS with MediatR")
p("What it is: CQRS (Command Query Responsibility Segregation) is the simple idea of separating "
  "actions that change data (commands, e.g. “approve this claim”) from actions that only read "
  "data (queries, e.g. “list my payslips”). MediatR is a small library that routes each request "
  "to exactly one handler, so the web controllers stay thin and every use-case lives in its own "
  "small, named class.")
p("Why we chose it: it makes the codebase predictable — one file per use-case, easy to find, easy "
  "to test in isolation. It also gives us a “pipeline” where cross-cutting behaviour (validation, "
  "logging, auditing) can be added once and applied to every request automatically.")
p("What we did not use, and why:")
bullets([
    ("Fat service classes", "A single “ClaimService” with twenty methods becomes a dumping ground that is hard to navigate and test."),
    ("Logic inside controllers", "Mixes web concerns with business logic and cannot be reused or tested without the web layer."),
])

h2("4.4  Input validation: FluentValidation")
p("What it is: a library for expressing rules about incoming data (“amount must be greater than "
  "zero”, “e-mail must be unique”) in clear, readable classes.")
p("Why we chose it: validation lives next to each use-case, reads almost like English, and is run "
  "automatically for every command via the MediatR pipeline — so invalid data is rejected with a "
  "clean error message before it ever reaches the business logic.")
p("What we did not use, and why: the framework’s built-in data-annotation attributes "
  "(e.g. [Required]) are fine for trivial checks but become awkward for rules that need the database "
  "(like uniqueness) or that depend on more than one field.")

h2("4.5  Data access: Entity Framework Core, with SQLite locally and Azure SQL in the cloud")
p("What it is: Entity Framework Core (EF Core) is an Object-Relational Mapper (ORM) — it lets the "
  "code work with normal C# objects while EF Core handles reading and writing the database. SQLite "
  "is a tiny, file-based database that needs no installation, ideal for development. Azure SQL is "
  "Microsoft’s fully-managed cloud database.")
p("Why we chose it: EF Core lets us write the data logic once and switch the underlying database by "
  "changing one line of configuration — SQLite for a zero-install developer laptop, Azure SQL in "
  "production. “Migrations” (versioned database changes) keep the schema in step with the code.")
p("What we did not use, and why:")
bullets([
    ("Dapper or raw SQL", "Faster for hand-tuned queries, but far more boilerplate and no automatic schema management; EF Core’s productivity wins for a business app of this size."),
    ("A NoSQL database (e.g. MongoDB / Cosmos DB)", "Our data is highly relational (users, claims, approvals, components all reference each other), which is exactly what a relational database does best. NoSQL shines for different shapes of data."),
    ("Installing SQL Server locally for development", "Heavyweight; SQLite gives the same EF Core experience with zero installation, supporting the “zero-install Codespaces” goal."),
])

h2("4.6  Security: JWT tokens, role-based access, and BCrypt password hashing")
p("What it is: when a user logs in, the server issues a JWT — a digitally-signed token the browser "
  "sends with each request to prove who it is. Each token carries the user’s role, and endpoints "
  "are guarded by role (RBAC = Role-Based Access Control). Passwords are never stored directly; they "
  "are run through BCrypt, a deliberately slow hashing algorithm that makes stolen passwords "
  "extremely hard to crack.")
p("Why we chose it: JWTs are stateless (the server doesn’t need to remember sessions), which scales "
  "well and works cleanly across the four separate frontends. RBAC matches the business exactly — "
  "an Employee simply cannot reach a Finance action. BCrypt is the long-standing, well-understood "
  "standard for password storage.")
p("What we did not use, and why:")
bullets([
    ("Server-side session cookies", "Require the server to store session state, which complicates scaling and cross-application sharing between the portals."),
    ("Storing passwords with fast hashes (MD5/SHA-256) or, worse, plain text", "Fast hashes are crackable at scale; plain text is a critical security failure. BCrypt’s slowness is a feature."),
])
p("In production this is intended to move to Microsoft Entra ID (formerly Azure AD) so the company’s "
  "existing corporate logins and single sign-on can be reused — see Section 6.")

h2("4.7  Compensation: the NRules business-rules engine")
p("What it is: salary is built from components (basic pay, house-rent allowance, bonus, provident-"
  "fund deduction, professional tax, and so on), and the formulas depend on grade and policy. NRules "
  "is a “rules engine” — each rule (“HRA is 40% of basic”, “bonus depends on grade band”) is "
  "written as its own small, declarative unit, and the engine runs them all to produce the result.")
p("Why we chose it: compensation policy changes often and is exactly the kind of logic that turns "
  "into tangled if-then-else code. Expressing each rule separately keeps it readable and lets new "
  "rules be added without rewriting existing ones — the rules become data-like and auditable.")
p("What we did not use, and why: hand-coded if/else or a big calculation method would work for "
  "today’s rules but rots quickly as policies multiply and interact. A rules engine is the "
  "purpose-built tool for “many independent business rules”.")

h2("4.8  Claims: an explicit state machine")
p("What it is: an expense claim moves through defined stages — Submitted → Under Review → "
  "Approved or Rejected → Settled. A “state machine” encodes exactly which moves are allowed (you "
  "cannot settle a claim that was never approved) and records every transition.")
p("Why we chose it: approval workflows are full of subtle rules about valid next steps. Making the "
  "states and allowed transitions explicit prevents illegal jumps, produces a natural audit trail, "
  "and turns invalid actions into clear ‘409 Conflict’ errors rather than silent data corruption.")
p("What we did not use, and why: a simple “status” text column with scattered checks is easy to "
  "get wrong — nothing stops the code from setting an impossible status. The state machine centralises "
  "and guarantees the rules.")

h2("4.9  Documents: file storage plus OCR (Tesseract locally, Azure AI Document Intelligence in the cloud)")
p("What it is: when a claim has a receipt, the file is saved through a storage abstraction, and OCR "
  "(Optical Character Recognition) reads the text from the image so, for example, the total can be "
  "extracted automatically. Locally this uses Tesseract, a free open-source OCR engine; in the cloud "
  "it is designed to use Azure AI Document Intelligence, which is far better at understanding receipts "
  "and invoices.")
p("Why we chose it: storing files and reading them are classic ‘external service’ concerns, so they "
  "sit behind interfaces (seams). That means the local engine can be a simple stub or Tesseract, while "
  "production swaps in Azure’s managed, highly-accurate service — without changing any business code.")
p("What we did not use, and why: building OCR ourselves is infeasible; a generic image library cannot "
  "read text. Tesseract is the standard free option for local use, and Azure AI Document Intelligence "
  "is purpose-built for documents in production.")

h2("4.10  Payroll: Polly for resilience, plus an asynchronous background worker")
p("What it is: paying money out means talking to an external payroll system that can be slow or "
  "briefly fail. Polly is a resilience library that automatically retries transient failures, can "
  "‘trip a circuit breaker’ if the other system is clearly down, and enforces timeouts. Settlements "
  "are also processed asynchronously: the request is accepted immediately, placed on an in-memory "
  "queue, and a background worker does the actual push so the user is never left waiting.")
p("Why we chose it: external systems are unreliable; without retries a momentary blip would fail a "
  "payment. Polly turns flaky calls into robust ones. Doing the work in the background keeps the user "
  "interface responsive and models how real payment processing behaves (‘accepted, processing…, "
  "done’).")
p("What we did not use, and why:")
bullets([
    ("Naive direct calls with no retry", "A single network hiccup would wrongly fail a settlement."),
    ("Hand-rolled retry loops", "Easy to get wrong (no back-off, no circuit breaker); Polly is the battle-tested standard."),
    ("Making the user wait for the external system (synchronous)", "Poor experience and fragile; asynchronous processing is the correct model for slow, external work."),
])
p("In production the in-memory queue is intended to become Azure Service Bus so messages survive "
  "restarts and can be processed by separate workers — see Section 6.")

h2("4.11  Promotions: event-driven notifications (MediatR locally, Azure Service Bus in the cloud)")
p("What it is: when a promotion is approved, the system ‘publishes an event’ (‘PromotionApproved’). "
  "Other parts of the system can ‘subscribe’ and react — here, an e-mail is sent to congratulate the "
  "employee — without the promotion logic needing to know about them.")
p("Why we chose it: this keeps features decoupled. The promotion code’s job is to approve a "
  "promotion; sending e-mail is a separate concern that simply reacts to the event. New reactions "
  "(update a dashboard, notify a manager) can be added later without touching the promotion code.")
p("What we did not use, and why: calling the e-mail service directly from the promotion logic would "
  "hard-wire the two together and make every new side-effect a change to core logic. Events keep the "
  "core focused and the reactions pluggable.")

h2("4.12  Reporting & Audit: LINQ aggregation, ClosedXML for Excel, and an audit pipeline")
p("What it is: reporting uses LINQ (a built-in C# way to query and summarise data) to produce "
  "dashboards — claims by status, settlements, headcount, and so on. ClosedXML is a library that "
  "generates real Excel (.xlsx) files for download. Separately, an ‘audit pipeline’ automatically "
  "records every action (who did what, and whether it succeeded) for accountability.")
p("Why we chose it: LINQ keeps summarising logic in readable C# rather than scattered SQL. Finance "
  "and HR live in Excel, so exporting real spreadsheets is genuinely useful. The audit log is "
  "implemented once as a pipeline step and therefore captures every command across all seven modules "
  "with no extra code per feature.")
p("What we did not use, and why:")
bullets([
    ("A heavyweight reporting/BI tool (e.g. Power BI) for this scope", "Excellent for rich analytics, but overkill for in-app operational reports and an export button; it would add cost and integration effort. (It remains a great future addition on Azure.)"),
    ("Adding audit code by hand in every handler", "Repetitive and easy to forget; the pipeline approach guarantees consistent coverage."),
])

h2("4.13  Logging: Serilog")
p("What it is: Serilog is a structured-logging library — it records what the application is doing "
  "in a queryable, structured way rather than as plain text.")
p("Why we chose it: good logs are essential for understanding and debugging a running system, and "
  "Serilog plugs neatly into Azure Application Insights in production for centralised monitoring.")

h2("4.14  The key idea that ties it together: integration ‘seams’")
p("Across the modules you will notice a recurring pattern: anything that touches the outside world "
  "(e-mail, payroll, file storage, OCR, the event bus) is hidden behind an interface — a contract "
  "that says what the capability does, not how. The business code depends only on the contract. We "
  "then provide two implementations: a simple local one for development and an Azure one for "
  "production, and choose between them with configuration.")
p("This single idea is what makes the whole ‘runs on a laptop with zero installation, yet deploys "
  "to Azure unchanged’ goal possible. It is dependency inversion applied consistently, and it is the "
  "backbone of the architecture.")

# =====================================================================
# 5. FRONTEND DECISIONS
# =====================================================================
h1("5. Frontend technology decisions")

h2("5.1  React 18")
p("What it is: React is the most widely-used library for building interactive web interfaces out of "
  "reusable ‘components’.")
p("Why we chose it: it is the industry standard with the largest ecosystem and talent pool, it pairs "
  "naturally with TypeScript, and — crucially — it is fully supported by Module Federation, the "
  "micro-frontend technique central to this design.")
p("What we did not use, and why: Angular (heavier, more opinionated — a fine choice but less "
  "flexible for independently-built micro-frontends) and Vue (smaller enterprise/Azure ecosystem). "
  "React gave the best balance of ubiquity, flexibility, and federation support.")

h2("5.2  Webpack 5 Module Federation (micro-frontends)")
p("What it is: the four portals (Employee, Manager, HR, Finance) are built as separate applications. "
  "A ‘shell’ application then loads the right portal at run-time and shows it as one seamless site. "
  "Module Federation is the Webpack feature that makes one app load code from another, live.")
p("Why we chose it: it matches the business — four audiences, four bodies of screens — and lets each "
  "portal be developed, built, and (in principle) deployed independently while sharing one login and "
  "look-and-feel. It is the frontend equivalent of the backend’s modular design.")
p("What we did not use, and why:")
bullets([
    ("One single large React app (a ‘monolithic SPA’)", "Simplest, but everything is coupled and deployed together — it would not demonstrate the independent-portal architecture the case study calls for."),
    ("iframes to embed separate apps", "Crude, with awkward sizing, navigation, and sharing of login; Module Federation integrates far more cleanly."),
    ("single-spa or Nx", "Capable micro-frontend frameworks, but Webpack’s native Module Federation is the most direct, widely-understood way to achieve this and needs no extra orchestration framework."),
])

h2("5.3  TypeScript")
p("What it is: TypeScript is JavaScript with types added — the editor and compiler catch many "
  "mistakes (wrong field names, wrong shapes) before the code runs.")
p("Why we chose it: it matches the strongly-typed backend, makes the data contracts between frontend "
  "and API explicit, and greatly improves safety and editor help on a multi-app codebase.")
p("What we did not use, and why: plain JavaScript is quicker to start but loses all of that safety — "
  "a poor trade-off on a system with this many moving parts.")

h2("5.4  An npm-workspaces monorepo")
p("What it is: all five frontend apps (shell + four portals) plus a shared library live in one "
  "repository, managed together by npm ‘workspaces’.")
p("Why we chose it: common code (the API client, login handling, UI building blocks, shared types) "
  "is written once in a ‘shared’ package and reused by every portal, and a single install sets up "
  "everything. It keeps the portals consistent and avoids copy-paste.")
p("What we did not use, and why: five separate repositories would duplicate the shared code and make "
  "coordinated changes painful for a single team.")

h2("5.5  Babel (not ts-loader) for building")
p("What it is: a small build-tool choice — Babel transpiles the TypeScript/React code to browser "
  "JavaScript quickly.")
p("Why we chose it: it is fast and simple for development. (Type-checking can still be run separately "
  "by the editor and the TypeScript compiler, so safety is not lost.)")

h2("5.6  Hand-written CSS design system (not a heavy UI library)")
p("What it is: the visual style (cards, buttons, tables, badges) is a small, shared set of components "
  "and one stylesheet.")
p("Why we chose it: it keeps the apps lightweight and fully under our control, with a consistent look "
  "across portals and no large dependency to manage.")
p("What we did not use, and why: large component libraries (Material UI, Ant Design, Bootstrap) are "
  "excellent for big teams but add weight and opinionated styling we did not need for a focused case "
  "study; a small bespoke system was leaner and clearer.")

h2("5.7  Sharing the login across portals via the browser’s localStorage")
p("What it is: after login, the token is kept in the browser’s localStorage, which every portal can "
  "read.")
p("Why we chose it: because the portals are technically separate applications, a shared browser store "
  "is the simplest reliable way for all of them to know the user is logged in, without a complex "
  "shared-memory mechanism.")

# =====================================================================
# 6. LOCAL VS AZURE
# =====================================================================
h1("6. Local versus Azure: component-by-component")
p("The architecture’s ‘seams’ (Section 4.14) mean the same code runs two ways: a zero-install "
  "set-up on a developer’s machine, and a managed, scalable set-up on Microsoft Azure. The table "
  "below maps each concern to both, and explains the Azure choice.")
spacer()
figure("as_built_deployment.png",
       "Figure 2 — The same system in two environments. Each row is one concern; only the implementation "
       "behind the seam changes between local development and Microsoft Azure. The seven business modules "
       "are identical in both.", width=5.8)
spacer()
table(
    ["Concern", "Local (development)", "Azure (production)", "Why the Azure service"],
    [
        ["Backend hosting", "Kestrel web server on the laptop", "Azure App Service or Azure Container Apps", "Managed, auto-scaling hosting with easy deployment and SSL."],
        ["Frontend hosting", "Webpack dev servers", "Azure Static Web Apps (or App Service)", "Purpose-built for hosting SPA/static front-ends with global delivery."],
        ["Database", "SQLite (single file)", "Azure SQL Database", "Fully-managed relational database; EF Core switches with one config line."],
        ["Authentication", "Self-issued JWT tokens", "Microsoft Entra ID (Azure AD) / Entra External ID", "Reuses corporate single sign-on; no password handling of our own."],
        ["Secrets (keys, connection strings)", "Local config files", "Azure Key Vault", "Secrets are centralised, encrypted, and never live in code."],
        ["E-mail", "Logged to console / smtp4dev", "Azure Communication Services (Email)", "Managed, scalable transactional e-mail."],
        ["File / receipt storage", "Local file system", "Azure Blob Storage", "Cheap, durable, scalable object storage for documents."],
        ["OCR (reading receipts)", "Tesseract / stub", "Azure AI Document Intelligence", "Pre-trained for receipts/invoices; far higher accuracy with no ML work."],
        ["Async messaging / events", "In-memory channel + MediatR", "Azure Service Bus", "Durable queues/topics that survive restarts and decouple workers."],
        ["Background processing", "Hosted background worker", "Azure Functions / Container Apps job", "Independent, elastic processing of queued work."],
        ["Resilience (retries)", "Polly (in-process)", "Polly + Azure platform retries", "Polly stays; Azure adds infrastructure-level resilience."],
        ["Logging & monitoring", "Serilog to console", "Application Insights (via Serilog)", "Centralised telemetry, dashboards, and alerting."],
        ["Reporting/analytics (future)", "In-app LINQ + Excel export", "Power BI / Azure Synapse (optional)", "For richer, organisation-wide analytics if needed later."],
    ],
    widths=[1.4, 1.6, 1.7, 2.1],
)
p("The important takeaway: moving from local to Azure is mostly a matter of configuration and "
  "swapping the implementation behind each seam — the business logic in the seven modules does not "
  "change.")

# =====================================================================
# 7. CONSOLIDATED ALTERNATIVES
# =====================================================================
h1("7. Quick reference: what we chose vs. what we rejected")
table(
    ["Area", "Chosen", "Main alternatives rejected", "Reason in one line"],
    [
        ["Backend language", ".NET 8 / C#", "Java, Node.js, Python", "Best Azure fit + strong typing for complex rules."],
        ["Architecture", "Clean modular monolith", "Microservices, plain N-tier", "Separation without distributed-systems overhead."],
        ["Request handling", "CQRS + MediatR", "Fat services, controller logic", "One small, testable class per use-case."],
        ["Validation", "FluentValidation", "Data-annotation attributes", "Readable, DB-aware, automatic rules."],
        ["Data access", "EF Core", "Dapper/raw SQL, NoSQL", "Productivity + relational data + DB portability."],
        ["Database", "SQLite → Azure SQL", "Local SQL Server, MongoDB", "Zero-install dev; managed relational in prod."],
        ["Passwords", "BCrypt", "MD5/SHA, plain text", "Deliberately slow = hard to crack."],
        ["Auth", "JWT + RBAC → Entra ID", "Session cookies", "Stateless, scalable, SSO-ready."],
        ["Compensation", "NRules engine", "Hand-coded if/else", "Many evolving rules stay readable."],
        ["Claims workflow", "State machine", "Status column + ad-hoc checks", "Prevents illegal transitions."],
        ["OCR", "Tesseract → Doc Intelligence", "Build our own", "Use proven engines via a seam."],
        ["Payroll calls", "Polly + async worker", "Direct synchronous calls", "Resilient + responsive."],
        ["Events", "MediatR → Service Bus", "Direct service calls", "Decoupled, pluggable reactions."],
        ["Excel reports", "ClosedXML", "Power BI for everything", "Right-sized in-app export."],
        ["Frontend", "React + Module Federation", "Single SPA, iframes, single-spa", "Independent portals, cleanly integrated."],
        ["Frontend language", "TypeScript", "Plain JavaScript", "Type safety across many apps."],
        ["UI styling", "Bespoke CSS system", "Material UI / Bootstrap", "Lightweight, consistent, controllable."],
    ],
    widths=[1.3, 1.6, 1.8, 2.1],
)

# =====================================================================
# 8. CROSS-CUTTING
# =====================================================================
h1("8. Cross-cutting concerns (applied everywhere)")
bullets([
    ("Security", "Every action is guarded by role; passwords are hashed; tokens are signed; secrets move to Key Vault in production."),
    ("Validation", "Every command is validated automatically before it runs, returning clear error messages."),
    ("Consistent errors", "A single error-handling layer turns problems into predictable web responses (e.g. 400 invalid, 401 unauthorised, 404 not found, 409 conflict)."),
    ("Auditing", "Every change is recorded once, centrally, for accountability."),
    ("Observability", "Structured logging throughout, ready for Application Insights in Azure."),
    ("Portability (local ↔ Azure)", "External services sit behind seams so the environment can change without touching business logic."),
])

# =====================================================================
# 9. SUMMARY
# =====================================================================
h1("9. Summary: the guiding principles")
p("If you remember only a few ideas from this document, make them these:")
numbered([
    "Organise around the business. The seven modules mirror real business capabilities, each cohesive and as independent as possible.",
    "Protect the business logic. Clean Architecture keeps the valuable rules at the centre, independent of database, web, and external services.",
    "Hide the outside world behind seams. E-mail, payroll, storage, OCR, and messaging are contracts with swappable implementations — this is what makes ‘zero-install locally, Azure in production’ possible.",
    "Use the right specialised tool for each hard problem. A rules engine for compensation, a state machine for approvals, a resilience library for unreliable calls, events for decoupled reactions.",
    "Right-size every choice. Prefer the simplest option that meets the need (modular monolith over microservices, bespoke CSS over a heavy UI kit, in-app reports over a BI platform) — while leaving a clear path to scale up on Azure.",
])
p("The result is a system that is easy to understand, safe to change, pleasant to develop locally, "
  "and ready to run on Microsoft Azure — without rewriting the parts that capture the business.")

# =====================================================================
# GLOSSARY
# =====================================================================
h1("Glossary (plain-language definitions)")
table(
    ["Term", "What it means"],
    [
        ["API", "The backend’s set of web endpoints that the frontends call to do things and fetch data."],
        ["Backend / Frontend", "Backend = the server-side engine and rules. Frontend = what runs in the browser."],
        ["Module", "A self-contained slice of the system for one business area (e.g. Claims)."],
        ["Clean Architecture", "A way of layering code so business rules don’t depend on technical details."],
        ["Modular monolith", "One deployable app that is cleanly divided into modules inside."],
        ["Microservices", "An alternative where each capability is its own separately-deployed service."],
        ["CQRS", "Separating actions that change data (commands) from those that read data (queries)."],
        ["MediatR", "A library that sends each request to its single handler and supports a pipeline."],
        ["ORM / EF Core", "A tool that maps database rows to code objects (EF Core is Microsoft’s)."],
        ["Migration", "A versioned, repeatable change to the database structure."],
        ["JWT", "A signed token proving who a logged-in user is, sent with each request."],
        ["RBAC", "Role-Based Access Control — permissions decided by the user’s role."],
        ["BCrypt", "A deliberately slow algorithm for safely storing passwords."],
        ["Rules engine (NRules)", "A tool that runs many independent business rules to compute a result."],
        ["State machine", "A model of allowed stages and the legal moves between them."],
        ["OCR", "Optical Character Recognition — reading text out of an image."],
        ["Polly", "A library that retries and protects calls to unreliable external systems."],
        ["Asynchronous", "Work accepted now and finished in the background, so users don’t wait."],
        ["Event bus", "A way to announce that something happened so other code can react."],
        ["Seam / abstraction", "A contract hiding a capability, with swappable implementations."],
        ["Module Federation", "A Webpack feature that loads one app’s code into another at run-time."],
        ["Micro-frontend", "A frontend split into independently-built apps combined into one site."],
        ["Monorepo", "One repository holding several related applications/packages."],
        ["TypeScript", "JavaScript with types, catching mistakes before the code runs."],
        ["LINQ", "A built-in C# way to query and summarise collections of data."],
        ["Azure", "Microsoft’s cloud platform that hosts the production version of the system."],
    ],
    widths=[1.7, 5.1],
)

import os
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, "Unified_Rewards_Platform_Architecture_Decisions.docx")
doc.save(out_path)
print("SAVED:", out_path)
