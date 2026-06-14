# -*- coding: utf-8 -*-
"""Generates (a) a target production deployment diagram and (b) the Production-Readiness /
Non-Functional-Requirements (NFR) Hardening Plan Word document for the Unified Rewards Platform,
written so a newcomer can follow it but actionable for engineers."""

import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.dirname(os.path.abspath(__file__))
INK = "#0f172a"
ACCENT = RGBColor(0x31, 0x2E, 0x81)
HEADER_FILL = "312E81"

# =====================================================================
# TARGET PRODUCTION DEPLOYMENT DIAGRAM
# =====================================================================
PAL = {
    "user":  ("#EEF2FF", "#6366F1"),
    "edge":  ("#E0E7FF", "#4F46E5"),
    "fe":    ("#DBEAFE", "#2563EB"),
    "api":   ("#DCFCE7", "#16A34A"),
    "inst":  ("#FFFFFF", "#16A34A"),
    "data":  ("#FEF3C7", "#D97706"),
    "work":  ("#FCE7F3", "#DB2777"),
    "cross": ("#E2E8F0", "#475569"),
    "band":  ("#F8FAFC", "#CBD5E1"),
}


def dbox(ax, x, y, w, h, text, kind, fs=8.5, bold=False):
    fill, edge = PAL[kind]
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.12",
                                lw=1.4, edgecolor=edge, facecolor=fill))
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs,
            color=INK, weight="bold" if bold else "normal")


def darrow(ax, x1, y1, x2, y2, text=None, color="#475569"):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=1.5, shrinkA=2, shrinkB=2))
    if text:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2, text, ha="center", va="center", fontsize=7.5,
                color=color, backgroundcolor="white")


def dband(ax, x, y, w, h, label):
    fill, edge = PAL["band"]
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.2",
                                lw=1.4, edgecolor=edge, facecolor=fill, linestyle="--"))
    ax.text(x + 0.8, y + h - 0.8, label, ha="left", va="top", fontsize=9.5, color="#334155", weight="bold")


fig, ax = plt.subplots(figsize=(13.5, 10.5))
ax.set_xlim(0, 100); ax.set_ylim(0, 100); ax.axis("off")
ax.text(50, 98, "Unified Rewards Platform — Target Production Deployment on Azure (scaled)",
        ha="center", va="center", fontsize=14.5, weight="bold", color=INK)

dbox(ax, 33, 91, 34, 5, "Users — employees · managers · HR · finance (global)", "user", bold=True)
darrow(ax, 50, 91, 50, 88)
dbox(ax, 28, 82.5, 44, 5, "Azure Front Door + CDN  —  global entry · TLS · edge caching · WAF", "edge", bold=True)

# Frontend branch
darrow(ax, 38, 82.5, 25, 75.5)
dbox(ax, 5, 70, 40, 5.5, "Azure Static Web Apps\nShell host + 4 portal remotes (Module Federation)", "fe", fs=8)

# API branch
darrow(ax, 62, 82.5, 72, 75.5)
dbox(ax, 55, 70, 40, 5.5, "Azure API Management\nrate limiting · gateway · routing", "edge", fs=8, bold=True)
darrow(ax, 75, 70, 60, 64)
darrow(ax, 25, 70, 45, 64)
dbox(ax, 22, 58, 56, 6, "Load balancer  /  ingress  —  spreads traffic across instances", "api", bold=True)

# Auto-scaling instances
dband(ax, 16, 40, 68, 16.5, "Auto-scaling API tier  ·  App Service / Container Apps  ·  modular monolith  ·  N replicas (e.g. 2 → 20+, scale by CPU/queue)")
for i in range(4):
    dbox(ax, 20 + i * 15.5, 42.5, 13, 9, ("API\ninstance " + str(i + 1)) if i < 3 else "…  +N", "inst", fs=8)
darrow(ax, 50, 58, 50, 56.5)
ax.text(86, 48, "scale\nout / in", ha="center", va="center", fontsize=8, color="#16A34A", style="italic")

darrow(ax, 50, 40, 50, 37)

# Data & async services band
dband(ax, 4, 14, 92, 22, "Backing services (each behind an application ‘seam’)")
dbox(ax, 6, 26, 20, 7, "Azure SQL\nprimary (write)", "data", fs=8, bold=True)
dbox(ax, 6, 16.5, 20, 7, "Azure SQL\nread replica (reporting)", "data", fs=8)
dbox(ax, 28, 26, 18, 7, "Azure Cache\nfor Redis (hot reads)", "data", fs=8)
dbox(ax, 28, 16.5, 18, 7, "Azure Blob\nStorage (receipts)", "data", fs=8)
dbox(ax, 48, 22, 20, 7, "Azure Service Bus\ndurable queue", "data", fs=8, bold=True)
dbox(ax, 70, 22, 24, 7, "Background workers / Functions\n(competing consumers: settlements, OCR)", "work", fs=7.6)
darrow(ax, 68, 25.5, 70, 25.5)

# Cross-cutting strip
dband(ax, 4, 2, 92, 9.5, "Cross-cutting platform services")
dbox(ax, 6, 3.5, 27, 6, "Microsoft Entra ID\nSSO · MFA · OIDC", "cross", fs=8)
dbox(ax, 36.5, 3.5, 27, 6, "Azure Key Vault\nsecrets · keys (Managed Identity)", "cross", fs=8)
dbox(ax, 67, 3.5, 27, 6, "Application Insights / Monitor\ntracing · metrics · alerts", "cross", fs=8)

fig.tight_layout()
fig.savefig(os.path.join(OUT, "target_production_architecture.png"), dpi=150, bbox_inches="tight")
plt.close(fig)

# =====================================================================
# WORD DOCUMENT
# =====================================================================
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)


def h1(t):
    p = doc.add_heading(t, 1)
    for r in p.runs: r.font.color.rgb = ACCENT


def h2(t):
    p = doc.add_heading(t, 2)
    for r in p.runs: r.font.color.rgb = ACCENT


def para(t, italic=False, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic; r.bold = bold; return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0] + " "); r.bold = True; p.add_run(it[1])
        else:
            p.add_run(it)


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def table(headers, rows, widths=None, fs=9):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(htext); r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(fs)
        shade(c, HEADER_FILL)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val)); r.font.size = Pt(fs)
    if widths:
        for i, w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width = Inches(w)
    doc.add_paragraph()


def figure(fname, caption, width=6.4):
    doc.add_picture(os.path.join(OUT, fname), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


# ---- Title ----
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Unified Rewards Platform"); r.bold = True; r.font.size = Pt(27); r.font.color.rgb = ACCENT
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Production-Readiness & Non-Functional Requirements (NFR) Hardening Plan"); r.font.size = Pt(15); r.italic = True
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run("Preparing the platform to run as a real-time application for an enterprise of CTS / TCS scale "
          "(hundreds of thousands of employees): load-balanced, auto-scaling, high-performance, secure, "
          "observable, and resilient — without breaking under load.").font.size = Pt(11)
doc.add_paragraph()
n = doc.add_paragraph(); n.alignment = WD_ALIGN_PARAGRAPH.CENTER
n.add_run("Cognizant — Upgrade to Architect, Phase 3 Case Study").italic = True
doc.add_page_break()

# ---- 1. Purpose ----
h1("1. Purpose & context")
para("This document is a practical plan to take the Unified Rewards Platform from a working, verified "
     "reference build to a system that can run in production for a very large enterprise — the kind with "
     "hundreds of thousands of employees and bursty peaks (payroll runs, benefit-enrolment windows, "
     "appraisal cycles). It assumes no prior knowledge: each item explains the current state, the risk at "
     "scale, the target solution on Microsoft Azure, and how we will know it is done (acceptance criteria).")
para("The honest headline:")
bullets([
    ("The architecture style is right.", " The platform is a stateless, fully-asynchronous, cleanly-"
     "modularised application. That foundation scales horizontally and does not need re-architecting."),
    ("The current implementation is not yet production-ready at this scale.", " It was tuned for zero-"
     "install local development. A specific set of changes (below) is required before it can serve "
     "enterprise load reliably. Most of them are hardening within the existing modular monolith — not a rewrite."),
])
para("Scale assumptions (to be confirmed by load testing): ~350,000 named users; tens of thousands of "
     "concurrent users at seasonal peaks; the application tier scaling from a small baseline to many "
     "instances on demand. All numeric targets in this document are proposed starting points to validate, "
     "not contractual figures.", italic=True)

# ---- 2. How to read ----
h1("2. How to read this plan — priority legend")
table(["Priority", "Meaning"],
      [["P0 — Critical", "Will break or fail at enterprise scale as currently coded. Must fix before production."],
       ["P1 — Performance", "Efficiency and speed improvements; needed for a good experience at scale."],
       ["P2 — Resilience", "Correctness and 'never breaks anywhere' robustness under load and failure."],
       ["P3 — Security", "Enterprise security & compliance; non-negotiable for sensitive HR/payroll data."],
       ["P4 — Operations", "Observability, deployment, and run-time operations at scale."]],
      widths=[1.6, 5.2])

# ---- 3. Current foundation ----
h1("3. What is already production-aligned")
para("These existing properties are exactly what a scalable system needs and should be preserved:")
bullets([
    ("Stateless API (JWT).", " No server-side session memory, so any instance can serve any request — the prerequisite for load balancing and horizontal scaling."),
    ("Fully asynchronous (async/await).", " A single instance handles many concurrent users efficiently while waiting on the database or network."),
    ("Clean Architecture + integration 'seams'.", " External dependencies sit behind interfaces, so infrastructure can be swapped (local → Azure) and hotspots extracted without rewrites."),
    ("CQRS, validation, RBAC, and an audit pipeline.", " Predictable structure, automatic input validation, role-based security, and a built-in audit trail."),
    ("Resilience already proven on the payroll integration (Polly).", " The retry / circuit-breaker / timeout pattern is in place and can be applied to every external call."),
])

# ---- 4. Target architecture ----
h1("4. Target production architecture (Azure)")
para("The diagram below shows the production shape: users arrive through a global edge (Front Door + CDN); "
     "the four portals are served as static content; API traffic passes through API Management (rate "
     "limiting) and a load balancer to an auto-scaling pool of identical API instances. Behind the seams "
     "sit managed Azure services, with heavy and slow work pushed onto a durable queue and processed by "
     "separate background workers.")
figure("target_production_architecture.png",
       "Figure 1 — Target production deployment. The stateless modular-monolith API runs as many "
       "auto-scaling instances behind a load balancer; the database, cache, storage, and durable "
       "messaging are managed Azure services; identity, secrets, and monitoring are cross-cutting.")
h2("Load balancing & auto-scaling, in plain terms")
bullets([
    ("Load balancing", " means many identical copies of the API run at once, and incoming requests are "
     "spread evenly across them. Because the API is stateless, this 'just works'."),
    ("Auto-scaling", " means Azure automatically adds more copies when traffic (CPU, request rate, or "
     "queue length) rises, and removes them when it falls — so the system stays fast at peak and cheap off-peak."),
    ("The real limit is the database, not the app tier.", " So the plan invests heavily in database scaling, "
     "caching, read replicas, and offloading heavy work — that is where capacity is won or lost."),
])

# ---- 5. Hardening backlog ----
h1("5. Prioritised hardening backlog")

h2("P0 — Critical (will break at scale as currently coded)")
table(["Item", "Current state", "Risk at scale", "Target solution", "Done when…"],
      [["Pagination on all list endpoints", "Endpoints return whole tables (audit caps at 500)",
        "A single call loads 100,000s of rows → timeouts & out-of-memory", "Server-side paging + filtering + sorting; keyset paging for large tables",
        "No endpoint can return an unbounded result set; verified by load test"],
       ["Reporting aggregation", "Dashboard pulls full tables into memory and sums in C# (a SQLite limitation)",
        "Instant out-of-memory / multi-minute queries on large data", "Aggregate inside Azure SQL; run reports off a read replica or materialized views",
        "Reports computed in the database; no full-table loads into app memory"],
       ["Async settlement queue", "In-memory Channel + in-process worker",
        "Not durable and not shared across instances → lost or duplicated work", "Azure Service Bus + competing-consumer workers + idempotency keys",
        "Messages survive restarts; exactly-once effect; many instances safe"],
       ["File storage", "Local file system", "Files saved on one instance are invisible to others",
        "Azure Blob Storage (seam already defined)", "Receipts readable from any instance"],
       ["Database engine", "SQLite (single file)", "Single-node; no concurrency at scale",
        "Azure SQL with retry-on-transient + connection pooling", "Running on Azure SQL with tuned pooling"],
       ["Schema migrations", "Auto-applied on every app startup", "Instances race to migrate; risky on a live DB",
        "Run migrations as a separate, gated deployment step", "App startup never alters the schema"]],
      widths=[1.3, 1.4, 1.5, 1.5, 1.4], fs=8)

h2("P1 — Performance & efficiency")
table(["Item", "Current", "Target solution"],
      [["Caching", "Every read hits the database", "Azure Cache for Redis for hot reads (plans, user directory, dashboards) with sensible TTLs and invalidation"],
       ["Audit write cost", "Extra SaveChanges per command (2× DB writes)", "Buffer/▶offload audit via Service Bus or a lightweight append store; batch writes"],
       ["EF Core read efficiency", "Some tracked reads; heavy Includes", "AsNoTracking everywhere for reads, DTO projections, split queries, compiled queries, AddDbContextPool"],
       ["HTTP efficiency", "No compression/caching headers", "Brotli/Gzip compression, output caching, ETags/Cache-Control, CDN for static assets"],
       ["Connection resilience", "Default", "EnableRetryOnFailure, tuned pool sizes, command timeouts"]],
      widths=[1.5, 2.2, 3.0], fs=9)

h2("P2 — Resilience & correctness (‘must not break anywhere’)")
bullets([
    ("Health checks", " — add liveness & readiness probes so the load balancer routes only to healthy instances and restarts sick ones."),
    ("Resilience on every external call", " — apply Polly timeouts/retries/circuit-breakers to email, OCR, blob, and database, not just payroll."),
    ("Optimistic concurrency", " — add rowversion concurrency tokens so two users acting on the same record cannot silently overwrite each other."),
    ("Transactional outbox", " — persist domain events (e.g. ‘promotion approved → email’) in the same transaction and publish via Service Bus, so events are never lost or double-sent."),
    ("Rate limiting / throttling", " — protect against spikes and abuse (built-in .NET rate limiter and/or API Management policies)."),
    ("Graceful shutdown", " — drain in-flight background work on scale-in/deploy so nothing is dropped."),
])

h2("P3 — Security & compliance (non-negotiable for HR/payroll data)")
bullets([
    ("Enterprise identity", " — replace self-issued symmetric JWT with Microsoft Entra ID (OIDC, asymmetric RS256, key rotation) for corporate SSO + MFA."),
    ("Secrets management", " — move signing keys and connection strings to Azure Key Vault, accessed via Managed Identity (no secrets in code/config)."),
    ("Data protection", " — TLS 1.2+ in transit, Transparent Data Encryption at rest, consider field-level encryption for salary data; enforce HTTPS and security headers (CSP, HSTS)."),
    ("Privacy & compliance", " — GDPR/PII handling, data-residency, immutable audit, and a data-retention/archival policy."),
    ("Frontend token safety", " — reconsider storing the token in localStorage (XSS exposure); evaluate httpOnly cookies / Backend-for-Frontend pattern and a strict Content-Security-Policy."),
    ("Lock down CORS", " — restrict to the real production origins."),
])

h2("P4 — Observability & operations")
bullets([
    ("Centralised telemetry", " — ship logs/metrics/traces to Application Insights (OpenTelemetry): per-endpoint latency, error rates, and distributed tracing with correlation IDs across the async workers."),
    ("Alerting & dashboards", " — alert on SLO breaches (latency, error rate, queue backlog)."),
    ("CI/CD with zero-downtime", " — blue-green or canary deployments; migrations as a pipeline gate."),
    ("Infrastructure as Code", " — Bicep/Terraform for repeatable, reviewable environments; explicit autoscale rules."),
    ("Data lifecycle", " — partition and archive ever-growing tables (audit, payslips) with a retention policy."),
    ("Disaster recovery", " — geo-replication and a tested failover plan if the RTO/RPO targets require it."),
])

h2("Frontend performance (page load & overall experience)")
bullets([
    ("Production build optimisation", " — minification, tree-shaking, content-hashed assets, Brotli, served via CDN over HTTP/2."),
    ("Client-side data caching", " — adopt React Query / SWR for caching, request de-duplication, and background refresh (large perceived-speed gain)."),
    ("Lazy loading", " — load each portal’s heavy sections on demand and preload the federation entry points."),
    ("Environment configuration", " — the API base URL must be configured per environment (currently hard-coded)."),
    ("Global readiness", " — localisation (i18n) and accessibility (WCAG 2.1 AA) for a multi-country workforce."),
])

# ---- 6. NFR targets ----
h1("6. Non-functional requirements — measurable targets")
para("Proposed, testable targets (validate and adjust during load testing):")
table(["Quality", "Target (proposed)", "How it is achieved"],
      [["Scalability", "~350k users; tens of thousands concurrent at peak; API tier scales 2 → 20+ instances",
        "Stateless API + load balancer + auto-scale; DB scale-up + read replica; Redis; async offload"],
       ["Performance (latency)", "Reads p95 < 300 ms; writes p95 < 800 ms (excl. external calls); first page load < 2.5 s, subsequent < 1 s",
        "Caching, pagination, query tuning, CDN, compression, client data-cache"],
       ["Throughput", "Sustain peak request rate with headroom; large reports run asynchronously",
        "Horizontal scale; Service Bus for heavy work; read replica for reporting"],
       ["Availability", "99.9% initially (multi-AZ), path to 99.95%+",
        "Multiple instances/zones; health checks; managed services; graceful deploys"],
       ["Reliability / data integrity", "No lost or double-processed work; RPO ≤ 5 min; RTO ≤ 1 h",
        "Durable queue + idempotency + outbox; optimistic concurrency; geo-replication"],
       ["Security", "SSO+MFA; encryption in transit & at rest; OWASP Top-10 clean; pen-tested",
        "Entra ID, Key Vault, TDE/TLS, rate limiting, security headers"],
       ["Compliance", "GDPR/PII compliant; immutable audit; defined retention",
        "Audit pipeline, data classification, retention/archival policies"],
       ["Observability", "100% of requests traced; alert on SLO breach",
        "Application Insights + OpenTelemetry + dashboards/alerts"],
       ["Maintainability", "Fast, safe change; ≥ 70% test coverage of Application layer",
        "Modular monolith + clean seams + CQRS + CI/CD + IaC"],
       ["Accessibility & i18n", "WCAG 2.1 AA; multi-language",
        "Accessible components; localisation framework"],
       ["Cost efficiency", "Scale down off-peak; reduce DB load",
        "Auto-scale in/out; caching; right-sized tiers"]],
      widths=[1.5, 2.4, 2.9], fs=8.5)

# ---- 7. Roadmap ----
h1("7. Implementation roadmap (suggested sequencing)")
para("Ordered so that the things which fail first are fixed first, and each phase leaves the system in a "
     "shippable state.")
table(["Phase", "Focus", "Key work"],
      [["Phase 1", "P0 scale foundation", "Azure SQL + push aggregation to SQL; pagination everywhere; Blob storage; Service Bus + idempotent workers; migrations as a deploy step"],
       ["Phase 2", "P1 performance", "Redis caching; EF read tuning + DbContext pooling; compression/output caching; CDN; client data-cache (React Query)"],
       ["Phase 3", "P2 + P3 resilience & security", "Health checks; resilience on all seams; optimistic concurrency; outbox; rate limiting; Entra ID; Key Vault; data protection"],
       ["Phase 4", "P4 operations & validation", "App Insights + tracing + alerts; CI/CD blue-green; IaC + autoscale rules; load + chaos + security testing; DR drill"]],
      widths=[0.9, 1.8, 4.1], fs=9)

# ---- 8. Risk register ----
h1("8. Risk register — what breaks first if ignored")
table(["Risk", "Likelihood", "Impact", "Mitigation"],
      [["List endpoint returns 100,000s of rows", "High", "Severe (outage/OOM)", "Pagination (P0)"],
       ["Reporting query exhausts memory", "High", "Severe", "SQL-side aggregation + read replica (P0)"],
       ["Settlements lost/duplicated across instances", "High", "Severe (financial)", "Service Bus + idempotency (P0)"],
       ["Receipts unreadable on other instances", "High", "Major", "Blob Storage (P0)"],
       ["Concurrent edits overwrite each other", "Medium", "Major (data integrity)", "Optimistic concurrency (P2)"],
       ["Credential/secret exposure", "Medium", "Severe (security)", "Entra ID + Key Vault (P3)"],
       ["No visibility when slow/failing", "High", "Major (MTTR)", "App Insights + alerts (P4)"]],
      widths=[2.4, 1.0, 1.3, 2.1], fs=8.5)

# ---- 9. Validation ----
h1("9. Testing & validation strategy")
bullets([
    ("Load & soak testing", " — simulate peak concurrent users (e.g. payroll-run profile) to validate latency/throughput targets and tune autoscale."),
    ("Spike & stress testing", " — confirm graceful behaviour (rate limiting, queueing) beyond expected peaks rather than collapse."),
    ("Chaos / failure injection", " — kill instances, fail the DB/queue, and confirm health checks, retries, and circuit breakers behave."),
    ("Security testing", " — SAST/DAST, dependency scanning, and a penetration test before go-live."),
    ("Data-scale testing", " — run with production-like data volumes (millions of rows) to catch query and index issues early."),
])

# ---- 10. Summary ----
h1("10. Summary")
para("The Unified Rewards Platform is built on the correct architectural foundation for enterprise scale: "
     "stateless, asynchronous, cleanly modular, with swappable infrastructure seams. To run as a real-time "
     "application for a CTS/TCS-scale organisation it needs targeted hardening — most urgently the P0 items "
     "(pagination, database-side reporting, durable queuing, shared storage, managed database, and safe "
     "migrations), followed by performance, resilience, security, and operations work.")
para("Crucially, every item in this plan can be delivered within the modular monolith using managed Azure "
     "services — horizontal auto-scaling, Azure SQL with a read replica, Redis, Blob, Service Bus, Entra ID, "
     "Key Vault, and Application Insights. Microservices remain an optional, later, selective step for a "
     "specific proven hotspot, not a prerequisite for serving this scale.")

# ---- Glossary ----
h1("Glossary")
table(["Term", "Plain meaning"],
      [["Load balancing", "Spreading incoming requests evenly across many running copies of the app."],
       ["Horizontal scaling", "Adding more copies (instances) of the app to handle more users."],
       ["Auto-scaling", "Automatically adding/removing instances based on load."],
       ["Stateless", "The app keeps no per-user memory between requests, so any instance can serve anyone."],
       ["Throughput", "How many requests the system handles per second."],
       ["Latency / p95", "How long a request takes; p95 = 95% of requests are faster than this."],
       ["Read replica", "A read-only copy of the database used to offload reporting/queries."],
       ["Cache (Redis)", "A fast in-memory store for frequently-read data to avoid hitting the database."],
       ["Durable queue", "A reliable message line (Service Bus) that survives restarts."],
       ["Idempotency", "Doing the same operation twice has the same effect as once (prevents double-pay)."],
       ["Competing consumers", "Several workers reading one queue, each item handled by exactly one."],
       ["Optimistic concurrency", "Detecting and rejecting conflicting simultaneous edits."],
       ["Transactional outbox", "Saving an event with its data in one transaction so it is never lost."],
       ["Circuit breaker", "Temporarily stops calling a failing service to let it recover."],
       ["Health check", "An endpoint that reports whether an instance is alive and ready for traffic."],
       ["CDN", "A global network that serves static content from close to the user."],
       ["RPO / RTO", "Max acceptable data loss / downtime during a disaster."],
       ["NFR", "Non-functional requirement — how well the system performs, not what it does."]],
      widths=[1.8, 5.0], fs=9)

out = os.path.join(OUT, "Unified_Rewards_Platform_NFR_Hardening_Plan.docx")
doc.save(out)
print("SAVED:", out)
print("SAVED:", os.path.join(OUT, "target_production_architecture.png"))
